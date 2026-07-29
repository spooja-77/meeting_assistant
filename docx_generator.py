"""
docx_generator.py
-----------------
Builds a formatted Word (.docx) document containing the meeting summary,
minutes of meeting, action items, and key decisions.
"""

import os
import datetime
from docx import Document
from docx.shared import Pt
from config import DOCX_OUTPUT_DIR


def create_meeting_docx(title: str, insights: dict) -> str:
    """
    Generate a .docx file summarizing the meeting.

    Args:
        title: title to show at the top of the document.
        insights: dict with keys summary, mom, action_items, key_decisions.

    Returns:
        The filepath of the generated .docx file.
    """
    doc = Document()

    # --- Title ---
    heading = doc.add_heading(title, level=0)

    # --- Summary section ---
    doc.add_heading("Meeting Summary", level=1)
    doc.add_paragraph(insights.get("summary", "N/A"))

    # --- Minutes of Meeting section ---
    doc.add_heading("Minutes of Meeting (MoM)", level=1)
    doc.add_paragraph(insights.get("mom", "N/A"))

    # --- Action Items section (as a bulleted list) ---
    doc.add_heading("Action Items", level=1)
    action_items = insights.get("action_items", [])
    if action_items:
        for item in action_items:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No action items identified.")

    # --- Key Decisions section (as a bulleted list) ---
    doc.add_heading("Key Decisions", level=1)
    key_decisions = insights.get("key_decisions", [])
    if key_decisions:
        for decision in key_decisions:
            doc.add_paragraph(decision, style="List Bullet")
    else:
        doc.add_paragraph("No key decisions identified.")

    # --- Save to disk with a unique, timestamped filename ---
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "_")).strip()
    filename = f"{safe_title or 'meeting'}_{timestamp}.docx"
    filepath = os.path.join(DOCX_OUTPUT_DIR, filename)

    doc.save(filepath)
    return filepath
